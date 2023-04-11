from flask import request, jsonify
from app.lib.cosine_similarity_files import cosine_similarity_files
import base64
import PyPDF2
import io
import docx

from app import app

@app.route('/health')
def health():
    return "OK"

@app.route('/similarity', methods=['POST'])
def get_similarity():
    data = request.json
    print(data)
    doc1blob = data['doc1blob']
    print(doc1blob)
    doc1blob = bytes(base64.b64decode(doc1blob))
    doc2blob = data['doc2blob']
    print(doc2blob)
    doc2blob = bytes(base64.b64decode(doc2blob))
    if data['fileType'] == "pdf":
        pdf_reader = PyPDF2.PdfFileReader(io.BytesIO(doc1blob))
        doc1 = ""
        for i in range(pdf_reader.getNumPages()):
            page = pdf_reader.getPage(i)
            doc1 += page.extractText()
        pdf_reader = PyPDF2.PdfFileReader(io.BytesIO(doc2blob))
        doc2 = ""
        for i in range(pdf_reader.getNumPages()):
            page = pdf_reader.getPage(i)
            doc2 += page.extractText()
    elif data['fileType'] == "txt":
        with io.BytesIO(doc1blob) as file:
            doc1 = file.read()
        with io.BytesIO(doc2blob) as file:
            doc2 = file.read()
    elif data['fileType'] == "docx":
        docx1 = docx.Document(io.BytesIO(doc1blob))
        doc1 = '\n'.join([paragraph.text for paragraph in docx1.paragraphs])
        docx2 = docx.Document(io.BytesIO(doc2blob))
        doc2 = '\n'.join([paragraph.text for paragraph in docx2.paragraphs])

    similarity = cosine_similarity_files(str(doc1), str(doc2))

    threshold = 0.8

    if similarity > threshold:
        message = "The two documents are similar."
    else:
        message = "The two documents are not similar."
    
    response = {'similarity_score': similarity, 'message': message}
    print(response)
    # response = {'similarity_score': 0.99*100, 'message': "pl"}

    return jsonify(response)